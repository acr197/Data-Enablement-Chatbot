# # rig_app.py
# # ---------------------------------------------------------------------------
# # Section: imports
# # ---------------------------------------------------------------------------
# from __future__ import annotations
#
# import os
# import re
# import sys
# import json
# import time
# import shutil
# import hashlib
# import logging
# import traceback
# from dataclasses import dataclass, asdict
# from typing import Dict, List, Tuple, Optional, Any
#
# import numpy as np
# import faiss  # faiss-cpu
# import streamlit as st
# import fitz  # PyMuPDF
# from docx import Document as DocxDocument  # python-docx
#
# try:
#     import textract  # optional for .doc
#     TEXTRACT_OK = True
# except Exception:
#     TEXTRACT_OK = False
#
# from openai import OpenAI
#
# print("Imports ready.")
#
# # ---------------------------------------------------------------------------
# # Section: user config
# # ---------------------------------------------------------------------------
# from dotenv import load_dotenv
# load_dotenv()
#
# # Paths and defaults
# DEFAULT_PROJECT_ROOT: str = r"C:\Users\Andrew\PycharmProjects\Data Enablement Chatbot"
# DEFAULT_DOCS_DIR: str = os.getenv(
#     "DOCS_DIR",
#     os.path.join(DEFAULT_PROJECT_ROOT, "docs")
# )
# DEFAULT_INDEX_SUBDIR: str = "_rig_index"
#
# # Models and parameters
# CHAT_MODEL: str = "gpt-4o-mini"
# EMBED_MODEL: str = "text-embedding-3-large"  # 3072-dim
# EMBED_DIM: int = 3072
# TOP_K_DEFAULT: int = 6
# EMBED_BATCH_SIZE: int = 64
# GEN_TEMPERATURE: float = 0.2
# GEN_MAX_TOKENS: int = 800
#
# # Chunking
# CHUNK_TARGET_CHARS: int = 1800
# CHUNK_OVERLAP_CHARS: int = 250
#
# # Filenames
# FAISS_INDEX_NAME: str = "faiss.index"
# DOCSTORE_JSONL_NAME: str = "docstore.jsonl"
# MANIFEST_JSON_NAME: str = "manifest.json"
# IDS_JSON_NAME: str = "ids.json"
#
# # API key pulled from .env
# OPENAI_API_KEY: str = os.getenv("OPENAI_API_KEY")
# if not OPENAI_API_KEY:
#     raise RuntimeError("OPENAI_API_KEY is missing. Add it to your .env file.")
#
# print("User config ready.")
#
#
# # ---------------------------------------------------------------------------
# # Section: logging setup
# # ---------------------------------------------------------------------------
# LOG_PATH: str = os.path.join(DEFAULT_PROJECT_ROOT, "rig_app.log")
# os.makedirs(DEFAULT_PROJECT_ROOT, exist_ok=True)
# logging.basicConfig(
#     level=logging.INFO,
#     format="%(asctime)s %(levelname)s %(message)s",
#     handlers=[
#         logging.FileHandler(LOG_PATH, encoding="utf-8"),
#         logging.StreamHandler(sys.stdout),
#     ],
# )
# logger = logging.getLogger("rig_app")
# print("Logging setup ready.")
#
# # ---------------------------------------------------------------------------
# # Section: data models
# # ---------------------------------------------------------------------------
# @dataclass
# class ChunkMeta:
#     file_path: str
#     file_name: str
#     file_ext: str
#     file_type: str
#     page: Optional[int]  # 1-based for PDFs, else None
#     chunk_idx: int       # 1-based index within file
#     char_start: int
#     char_end: int
#     sha256: str
#
# @dataclass
# class ChunkRecord:
#     chunk_id: int        # int64 for FAISS
#     text: str
#     meta: ChunkMeta
#
# @dataclass
# class FileManifestEntry:
#     sha256: str
#     mtime: float
#     chunk_ids: List[int]
#
# print("Data models ready.")
#
# # ---------------------------------------------------------------------------
# # Section: utilities
# # ---------------------------------------------------------------------------
# CONTROL_CHAR_RE = re.compile(r"[\u0000-\u0008\u000B-\u001F\u007F\u200B\u200C\u200D\uFEFF]")
# TIMESTAMP_RE = re.compile(r"(\[?\b\d{1,2}:\d{2}:\d{2}\b\]?)")
# SPEAKER_ARROW_RE = re.compile(r"(?m)^(>{2,}\s*)")
# WHITESPACE_RUN_RE = re.compile(r"[ \t]{2,}")
# PUNCT_RUN_RE = re.compile(r"([!?.,]){3,}")
#
# def safe_mkdir(path: str) -> None:
#     try:
#         os.makedirs(path, exist_ok=True)
#     except Exception as e:
#         logger.error("Failed to create directory: %s | %s", path, e)
#
# def compute_sha256(path: str) -> str:
#     h = hashlib.sha256()
#     with open(path, "rb") as f:
#         for chunk in iter(lambda: f.read(1024 * 1024), b""):
#             h.update(chunk)
#     return h.hexdigest()
#
# def clean_text(text: str) -> str:
#     text = CONTROL_CHAR_RE.sub("", text)
#     text = TIMESTAMP_RE.sub("", text)
#     text = SPEAKER_ARROW_RE.sub("", text)
#     text = PUNCT_RUN_RE.sub(lambda m: m.group(1) * 2, text)
#     lines = [WHITESPACE_RUN_RE.sub(" ", ln) for ln in text.splitlines()]
#     return "\n".join(lines).strip()
#
# def paragraphs(text: str) -> List[str]:
#     parts = re.split(r"\n\s*\n", text)
#     return [p.strip() for p in parts if p.strip()]
#
# def chunk_text(text: str, target: int, overlap: int) -> List[Tuple[int, int, str]]:
#     """Return list of (start_char, end_char, chunk_text)."""
#     paras = paragraphs(text)
#     chunks: List[Tuple[int, int, str]] = []
#     cur: List[str] = []
#     start = 0
#     for p in paras:
#         if sum(len(x) + 2 for x in cur) + len(p) > target and cur:
#             chunk_str = "\n\n".join(cur)
#             end = start + len(chunk_str)
#             chunks.append((start, end, chunk_str))
#             if overlap > 0:
#                 keep = chunk_str[-overlap:]
#                 cur = [keep]
#                 start = end - len(keep)
#             else:
#                 cur = []
#                 start = end
#         cur.append(p)
#     if cur:
#         chunk_str = "\n\n".join(cur)
#         end = start + len(chunk_str)
#         chunks.append((start, end, chunk_str))
#     return chunks
#
# def stable_int64_id(s: str) -> int:
#     h = hashlib.sha256(s.encode("utf-8")).digest()
#     return int.from_bytes(h[:8], "big") & ((1 << 63) - 1)
#
# def ensure_index_dir(docs_dir: str) -> str:
#     idx_dir = os.path.join(docs_dir, DEFAULT_INDEX_SUBDIR)
#     safe_mkdir(idx_dir)
#     return idx_dir
#
# def read_json(path: str, default: dict) -> dict:
#     try:
#         if os.path.exists(path):
#             with open(path, "r", encoding="utf-8") as f:
#                 return json.load(f)
#     except Exception as e:
#         logger.error("Failed to read JSON %s: %s", path, e)
#     return default
#
# def write_json(path: str, data: dict) -> None:
#     try:
#         with open(path, "w", encoding="utf-8") as f:
#             json.dump(data, f, ensure_ascii=False, indent=2)
#     except Exception as e:
#         logger.error("Failed to write JSON %s: %s", path, e)
#
# def read_docstore_jsonl(path: str) -> Dict[int, ChunkRecord]:
#     records: Dict[int, ChunkRecord] = {}
#     try:
#         if os.path.exists(path):
#             with open(path, "r", encoding="utf-8") as f:
#                 for ln in f:
#                     if not ln.strip():
#                         continue
#                     obj = json.loads(ln)
#                     meta = ChunkMeta(**obj["meta"])
#                     rec = ChunkRecord(chunk_id=int(obj["chunk_id"]), text=obj["text"], meta=meta)
#                     records[rec.chunk_id] = rec
#     except Exception as e:
#         logger.error("Failed to read docstore %s: %s", path, e)
#     return records
#
# def write_docstore_jsonl(path: str, records: Dict[int, ChunkRecord]) -> None:
#     try:
#         with open(path, "w", encoding="utf-8") as f:
#             for rec in records.values():
#                 obj = {"chunk_id": rec.chunk_id, "text": rec.text, "meta": asdict(rec.meta)}
#                 f.write(json.dumps(obj, ensure_ascii=False) + "\n")
#     except Exception as e:
#         logger.error("Failed to write docstore %s: %s", path, e)
#
# print("Utilities ready.")
#
# # ---------------------------------------------------------------------------
# # Section: indexing
# # ---------------------------------------------------------------------------
# def load_or_create_faiss(index_path: str, dim: int) -> faiss.IndexIDMap2:
#     if os.path.exists(index_path):
#         try:
#             index = faiss.read_index(index_path)
#             if not isinstance(index, faiss.IndexIDMap2):
#                 index = faiss.IndexIDMap2(index)
#             return index  # type: ignore
#         except Exception as e:
#             logger.error("Failed to read FAISS index, recreating: %s", e)
#     base = faiss.IndexFlatIP(dim)
#     return faiss.IndexIDMap2(base)
#
# def save_faiss(index: faiss.IndexIDMap2, index_path: str) -> None:
#     try:
#         faiss.write_index(index, index_path)
#     except Exception as e:
#         logger.error("Failed to save FAISS index: %s", e)
#
# def remove_chunk_ids(index: faiss.IndexIDMap2, ids: List[int]) -> None:
#     if not ids:
#         return
#     try:
#         sel = faiss.IDSelectorBatch(np.array(ids, dtype=np.int64))
#         index.remove_ids(sel)
#     except Exception as e:
#         logger.error("Failed to remove ids: %s", e)
#
# def list_supported_files(docs_dir: str) -> List[str]:
#     paths: List[str] = []
#     for root, _, files in os.walk(docs_dir):
#         if os.path.basename(root) == DEFAULT_INDEX_SUBDIR:
#             continue
#         for fn in files:
#             ext = os.path.splitext(fn)[1].lower()
#             if ext in (".pdf", ".docx", ".txt", ".doc"):
#                 paths.append(os.path.join(root, fn))
#     return paths
#
# def incremental_plan(docs_dir: str, manifest_path: str) -> Tuple[Dict[str, FileManifestEntry], List[str], List[str]]:
#     manifest_raw = read_json(manifest_path, {})
#     manifest: Dict[str, FileManifestEntry] = {}
#     for k, v in manifest_raw.items():
#         try:
#             manifest[k] = FileManifestEntry(sha256=v["sha256"], mtime=float(v["mtime"]), chunk_ids=list(v.get("chunk_ids", [])))
#         except Exception:
#             continue
#     files = list_supported_files(docs_dir)
#     new_or_changed: List[str] = []
#     unchanged: List[str] = []
#     for path in files:
#         try:
#             mtime = os.path.getmtime(path)
#             sha = compute_sha256(path)
#             if path not in manifest or manifest[path].sha256 != sha or abs(manifest[path].mtime - mtime) > 1e-6:
#                 new_or_changed.append(path)
#             else:
#                 unchanged.append(path)
#         except Exception as e:
#             logger.error("Stat/hash failed for %s: %s", path, e)
#     return manifest, new_or_changed, unchanged
#
# print("Indexing scaffolding ready.")
#
# # ---------------------------------------------------------------------------
# # Section: extraction
# # ---------------------------------------------------------------------------
# def extract_pdf(path: str) -> List[Tuple[int, str]]:
#     out: List[Tuple[int, str]] = []
#     try:
#         with fitz.open(path) as doc:
#             for i, page in enumerate(doc, start=1):
#                 txt = page.get_text("text") or ""
#                 out.append((i, clean_text(txt)))
#     except Exception as e:
#         logger.error("PDF extract failed %s: %s", path, e)
#     return out
#
# def extract_docx(path: str) -> str:
#     try:
#         doc = DocxDocument(path)
#         parts: List[str] = []
#         for p in doc.paragraphs:
#             parts.append(p.text)
#         for tbl in doc.tables:
#             for row in tbl.rows:
#                 for cell in row.cells:
#                     parts.append(cell.text)
#         return clean_text("\n".join(parts))
#     except Exception as e:
#         logger.error("DOCX extract failed %s: %s", path, e)
#         return ""
#
# def extract_doc(path: str) -> str:
#     if not TEXTRACT_OK:
#         logger.warning(".doc support requires textract; skipping %s", path)
#         return ""
#     try:
#         raw = textract.process(path)
#         return clean_text(raw.decode("utf-8", errors="ignore"))
#     except Exception as e:
#         logger.error("DOC extract failed %s: %s", path, e)
#         return ""
#
# def extract_txt(path: str) -> str:
#     try:
#         with open(path, "r", encoding="utf-8", errors="ignore") as f:
#             return clean_text(f.read())
#     except Exception as e:
#         logger.error("TXT read failed %s: %s", path, e)
#         return ""
#
# def build_chunks_for_file(path: str) -> List[ChunkRecord]:
#     ext = os.path.splitext(path)[1].lower()
#     fname = os.path.basename(path)
#     sha = compute_sha256(path)
#     records: List[ChunkRecord] = []
#
#     if ext == ".pdf":
#         pages = extract_pdf(path)
#         for page_num, page_text in pages:
#             if not page_text.strip():
#                 continue
#             spans = chunk_text(page_text, CHUNK_TARGET_CHARS, CHUNK_OVERLAP_CHARS)
#             for idx, (start, end, chunk) in enumerate(spans, start=1):
#                 cid = stable_int64_id(f"{path}|{sha}|pdf|{page_num}|{idx}")
#                 meta = ChunkMeta(
#                     file_path=path,
#                     file_name=fname,
#                     file_ext=ext,
#                     file_type="pdf",
#                     page=page_num,
#                     chunk_idx=idx,
#                     char_start=start,
#                     char_end=end,
#                     sha256=sha,
#                 )
#                 records.append(ChunkRecord(chunk_id=cid, text=chunk, meta=meta))
#     elif ext == ".docx":
#         text = extract_docx(path)
#         spans = chunk_text(text, CHUNK_TARGET_CHARS, CHUNK_OVERLAP_CHARS)
#         for idx, (start, end, chunk) in enumerate(spans, start=1):
#             cid = stable_int64_id(f"{path}|{sha}|docx|{idx}")
#             meta = ChunkMeta(
#                 file_path=path,
#                 file_name=fname,
#                 file_ext=ext,
#                 file_type="docx",
#                 page=None,
#                 chunk_idx=idx,
#                 char_start=start,
#                 char_end=end,
#                 sha256=sha,
#             )
#             records.append(ChunkRecord(chunk_id=cid, text=chunk, meta=meta))
#     elif ext == ".doc":
#         text = extract_doc(path)
#         if text.strip():
#             spans = chunk_text(text, CHUNK_TARGET_CHARS, CHUNK_OVERLAP_CHARS)
#             for idx, (start, end, chunk) in enumerate(spans, start=1):
#                 cid = stable_int64_id(f"{path}|{sha}|doc|{idx}")
#                 meta = ChunkMeta(
#                     file_path=path,
#                     file_name=fname,
#                     file_ext=ext,
#                     file_type="doc",
#                     page=None,
#                     chunk_idx=idx,
#                     char_start=start,
#                     char_end=end,
#                     sha256=sha,
#                 )
#                 records.append(ChunkRecord(chunk_id=cid, text=chunk, meta=meta))
#     elif ext == ".txt":
#         text = extract_txt(path)
#         spans = chunk_text(text, CHUNK_TARGET_CHARS, CHUNK_OVERLAP_CHARS)
#         for idx, (start, end, chunk) in enumerate(spans, start=1):
#             cid = stable_int64_id(f"{path}|{sha}|txt|{idx}")
#             meta = ChunkMeta(
#                 file_path=path,
#                 file_name=fname,
#                 file_ext=ext,
#                 file_type="txt",
#                 page=None,
#                 chunk_idx=idx,
#                 char_start=start,
#                 char_end=end,
#                 sha256=sha,
#             )
#             records.append(ChunkRecord(chunk_id=cid, text=chunk, meta=meta))
#     else:
#         logger.info("Unsupported file skipped: %s", path)
#
#     return records
#
# print("Extraction ready.")
#
# # ---------------------------------------------------------------------------
# # Section: OpenAI helpers
# # ---------------------------------------------------------------------------
# def get_client() -> OpenAI:
#     try:
#         return OpenAI(api_key=OPENAI_API_KEY)
#     except Exception as e:
#         logger.error("OpenAI client init failed: %s", e)
#         raise
#
# def embed_texts(client: OpenAI, texts: List[str]) -> np.ndarray:
#     out: List[List[float]] = []
#     for i in range(0, len(texts), EMBED_BATCH_SIZE):
#         batch = texts[i:i + EMBED_BATCH_SIZE]
#         try:
#             resp = client.embeddings.create(model=EMBED_MODEL, input=batch)
#             for item in resp.data:
#                 out.append(item.embedding)
#         except Exception as e:
#             logger.error("Embedding batch failed: %s", e)
#             raise
#     arr = np.array(out, dtype=np.float32)
#     norms = np.linalg.norm(arr, axis=1, keepdims=True) + 1e-12
#     arr = arr / norms
#     return arr
#
# def chat_answer(client: OpenAI, system_prompt: str, user_prompt: str, temperature: float, max_tokens: int) -> str:
#     try:
#         resp = client.responses.create(
#             model=CHAT_MODEL,
#             temperature=temperature,
#             max_output_tokens=max_tokens,
#             input=[
#                 {"role": "system", "content": system_prompt},
#                 {"role": "user", "content": user_prompt},
#             ],
#         )
#         return resp.output_text  # type: ignore
#     except Exception as e:
#         logger.error("Chat failed: %s", e)
#         raise
#
# print("OpenAI helpers ready.")
#
# # ---------------------------------------------------------------------------
# # Section: retrieval+answer
# # ---------------------------------------------------------------------------
# def ensure_all_ready(docs_dir: str) -> Tuple[faiss.IndexIDMap2, Dict[int, ChunkRecord], Dict[str, FileManifestEntry], str, str, str, str]:
#     idx_dir = ensure_index_dir(docs_dir)
#     faiss_path = os.path.join(idx_dir, FAISS_INDEX_NAME)
#     docstore_path = os.path.join(idx_dir, DOCSTORE_JSONL_NAME)
#     manifest_path = os.path.join(idx_dir, MANIFEST_JSON_NAME)
#     ids_path = os.path.join(idx_dir, IDS_JSON_NAME)
#     index = load_or_create_faiss(faiss_path, EMBED_DIM)
#     docstore = read_docstore_jsonl(docstore_path)
#     manifest_raw = read_json(manifest_path, {})
#     manifest: Dict[str, FileManifestEntry] = {}
#     for k, v in manifest_raw.items():
#         try:
#             manifest[k] = FileManifestEntry(sha256=v["sha256"], mtime=float(v["mtime"]), chunk_ids=list(v.get("chunk_ids", [])))
#         except Exception:
#             continue
#     return index, docstore, manifest, faiss_path, docstore_path, manifest_path, ids_path
#
# def reindex(docs_dir: str, progress_cb=None) -> Tuple[int, int]:
#     index, docstore, manifest, faiss_path, docstore_path, manifest_path, _ = ensure_all_ready(docs_dir)
#     client = get_client()
#     start_time = time.time()
#
#     manifest, changed, unchanged = incremental_plan(docs_dir, manifest_path)
#     total_new_chunks = 0
#     total_removed_chunks = 0
#
#     for path in changed:
#         if path in manifest and manifest[path].chunk_ids:
#             remove_chunk_ids(index, manifest[path].chunk_ids)
#             total_removed_chunks += len(manifest[path].chunk_ids)
#             for cid in manifest[path].chunk_ids:
#                 if cid in docstore:
#                     del docstore[cid]
#
#     for i, path in enumerate(changed, start=1):
#         if progress_cb:
#             progress_cb(f"Processing {i}/{len(changed)}: {os.path.basename(path)}")
#         try:
#             records = build_chunks_for_file(path)
#             if not records:
#                 manifest[path] = FileManifestEntry(sha256=compute_sha256(path), mtime=os.path.getmtime(path), chunk_ids=[])
#                 continue
#             embeddings = embed_texts(client, [r.text for r in records])
#             ids = np.array([r.chunk_id for r in records], dtype=np.int64)
#             index.add_with_ids(embeddings, ids)
#             for rec in records:
#                 docstore[rec.chunk_id] = rec
#             manifest[path] = FileManifestEntry(
#                 sha256=compute_sha256(path),
#                 mtime=os.path.getmtime(path),
#                 chunk_ids=[r.chunk_id for r in records],
#             )
#             total_new_chunks += len(records)
#         except Exception as e:
#             logger.error("File index failed %s: %s", path, e)
#             logger.debug(traceback.format_exc())
#
#     save_faiss(index, faiss_path)
#     write_docstore_jsonl(docstore_path, docstore)
#     write_json(manifest_path, {k: asdict(v) for k, v in manifest.items()})
#
#     logger.info("Reindex done in %.2fs | new=%d removed=%d unchanged_files=%d", time.time() - start_time, total_new_chunks, total_removed_chunks, len(unchanged))
#     return total_new_chunks, total_removed_chunks
#
# def clear_index(docs_dir: str) -> None:
#     idx_dir = ensure_index_dir(docs_dir)
#     try:
#         shutil.rmtree(idx_dir, ignore_errors=True)
#     except Exception as e:
#         logger.error("Clear index failed: %s", e)
#     safe_mkdir(idx_dir)
#
# def format_citation(meta: ChunkMeta) -> str:
#     if meta.file_type == "pdf" and meta.page:
#         return f"[{meta.file_name} p. {meta.page}]"
#     return f"[{meta.file_name} chunk {meta.chunk_idx}]"
#
# def build_context_block(recs: List[ChunkRecord]) -> Tuple[str, List[str]]:
#     lines: List[str] = []
#     cits: List[str] = []
#     for rec in recs:
#         cit = format_citation(rec.meta)
#         cits.append(cit)
#         preview = rec.text.strip()
#         lines.append(f"{cit}\n{preview}")
#     return "\n\n---\n\n".join(lines), cits
#
# def search(index: faiss.IndexIDMap2, q_emb: np.ndarray, top_k: int) -> Tuple[np.ndarray, np.ndarray]:
#     if index.ntotal == 0:
#         return np.array([], dtype=np.int64), np.array([], dtype=np.float32)
#     D, I = index.search(q_emb, top_k)
#     return I[0], D[0]
#
# def generate_answer(client: OpenAI, query: str, candidates: List[ChunkRecord], temperature: float, max_tokens: int) -> str:
#     sys_prompt = "Use only the provided snippets. If insufficient, say so plainly."
#     context_block, _ = build_context_block(candidates)
#     user_prompt = (
#         "Question:\n"
#         f"{query}\n\n"
#         "Context snippets (use only these; cite each claim inline using the required format):\n"
#         f"{context_block}\n\n"
#         "Cite PDFs like [filename.pdf p. N]. Cite DOCX/TXT/DOC like [filename.ext chunk K]."
#     )
#     ans = chat_answer(client, sys_prompt, user_prompt, temperature, max_tokens)
#     if not re.search(r"\[[^\[\]]+\]", ans):
#         _, cits = build_context_block(candidates)
#         srcs = "\n".join(f"- {c}" for c in dict.fromkeys(cits))
#         ans = ans.rstrip() + "\n\nSources:\n" + srcs
#     return ans
#
# print("Retrieval+answer ready.")
#
# # ---------------------------------------------------------------------------
# # Section: Streamlit UI
# # ---------------------------------------------------------------------------
# def ui_write_requirements(project_root: str) -> str:
#     path = os.path.join(project_root, "requirements_rig.txt")
#     try:
#         with open(path, "w", encoding="utf-8") as f:
#             f.write(REQUIREMENTS_SNIPPET.strip() + "\n")
#         return path
#     except Exception as e:
#         logger.error("Write requirements failed: %s", e)
#         return ""
#
# def handle_uploads(target_dir: str, files: List[Any]) -> List[str]:
#     saved: List[str] = []
#     for uf in files:
#         try:
#             ext = os.path.splitext(uf.name)[1].lower()
#             if ext not in (".pdf", ".docx", ".txt", ".doc"):
#                 continue
#             dest = os.path.join(target_dir, uf.name)
#             with open(dest, "wb") as f:
#                 f.write(uf.getbuffer())
#             saved.append(dest)
#         except Exception as e:
#             logger.error("Upload save failed %s: %s", getattr(uf, 'name', 'unknown'), e)
#     return saved
#
# def main_app() -> None:
#     st.set_page_config(page_title="Local Docs RAG", layout="wide")
#     st.title("Local, Docs-Only RAG Chatbot")
#
#     st.sidebar.header("Controls")
#     docs_dir = st.sidebar.text_input("Documents folder", value=DEFAULT_DOCS_DIR)
#     if not docs_dir:
#         st.sidebar.error("Provide a documents folder.")
#         st.stop()
#     safe_mkdir(docs_dir)
#
#     with st.sidebar.expander("Upload files"):
#         files = st.file_uploader("Drop PDF/DOC/DOCX/TXT here", type=["pdf", "docx", "txt", "doc"], accept_multiple_files=True)
#         if files:
#             saved = handle_uploads(docs_dir, files)
#             if saved:
#                 st.success(f"Saved {len(saved)} file(s). Click Reindex.")
#
#     top_k = st.sidebar.slider("Top-k", min_value=1, max_value=12, value=TOP_K_DEFAULT, step=1)
#     temperature = st.sidebar.slider("Temperature", min_value=0.0, max_value=1.0, value=GEN_TEMPERATURE, step=0.05)
#     max_tokens = st.sidebar.slider("Max tokens", min_value=200, max_value=2000, value=GEN_MAX_TOKENS, step=50)
#     strict_mode = st.sidebar.checkbox("Strict: append Sources if no citations", value=True)
#     show_debug = st.sidebar.checkbox("Show debug info", value=False)
#
#     c1, c2, c3 = st.sidebar.columns(3)
#     if c1.button("Reindex"):
#         msg = st.sidebar.empty()
#         def cb(text: str) -> None:
#             msg.info(text)
#         newc, rmvc = reindex(docs_dir, progress_cb=cb)
#         st.sidebar.success(f"Reindex complete. New: {newc}, Removed: {rmvc}")
#     if c2.button("Clear index"):
#         clear_index(docs_dir)
#         st.sidebar.warning("Index cleared.")
#     if c3.button("Write requirements"):
#         p = ui_write_requirements(DEFAULT_PROJECT_ROOT)
#         if p:
#             st.sidebar.success(f"Wrote {p}")
#
#     q = st.text_area("Ask a question about your documents:", height=120, placeholder="Type your question...")
#     if st.button("Ask"):
#         if not q.strip():
#             st.error("Enter a question.")
#             st.stop()
#         index, docstore, _, _, _, _, _ = ensure_all_ready(docs_dir)
#         if index.ntotal == 0 or not docstore:
#             st.warning("No index found. Click Reindex after adding files.")
#             st.stop()
#
#         try:
#             client = get_client()
#             q_emb = embed_texts(client, [q])
#             I, D = search(index, q_emb, top_k)
#             if I.size == 0:
#                 st.write("I couldn’t find that in the provided documents.")
#                 st.stop()
#             records: List[ChunkRecord] = []
#             for cid in I.tolist():
#                 rec = docstore.get(int(cid))
#                 if rec:
#                     records.append(rec)
#             if not records:
#                 st.write("I couldn’t find that in the provided documents.")
#                 st.stop()
#
#             ans = generate_answer(client, q, records, temperature, max_tokens)
#             if strict_mode and not re.search(r"\[[^\[\]]+\]", ans):
#                 _, cits = build_context_block(records)
#                 srcs = "\n".join(f"- {c}" for c in dict.fromkeys(cits))
#                 ans = ans.rstrip() + "\n\nSources:\n" + srcs
#
#             st.markdown("**Answer**")
#             st.write(ans)
#
#             with st.expander("Top sources"):
#                 for rec in records:
#                     st.markdown(f"- {format_citation(rec.meta)}")
#                     st.code(rec.text[:1200] + ("..." if len(rec.text) > 1200 else ""))
#
#             if show_debug:
#                 st.subheader("Debug")
#                 st.json({
#                     "docs_dir": docs_dir,
#                     "index_size": int(index.ntotal),
#                     "retrieved_ids": [int(x) for x in I.tolist()],
#                     "scores": [float(x) for x in D.tolist()],
#                 })
#         except Exception as e:
#             st.error(f"Error: {e}")
#             logger.error("Ask failed: %s", e)
#             logger.debug(traceback.format_exc())
#
# print("Streamlit UI ready.")
#
# # ---------------------------------------------------------------------------
# # Section: requirements snippet
# # ---------------------------------------------------------------------------
# REQUIREMENTS_SNIPPET = """
# faiss-cpu==1.8.0.post1
# numpy==1.26.4
# PyMuPDF==1.24.9
# streamlit==1.36.0
# openai==1.51.0
# python-docx==1.1.2
# tiktoken==0.7.0
# textract==1.6.5
# pyinstaller==6.10.0
# """
#
# print("Requirements snippet ready.")
#
# # ---------------------------------------------------------------------------
# # Section: final notes
# # ---------------------------------------------------------------------------
# def _in_streamlit_runtime() -> bool:
#     """Return True only when executed by Streamlit, False when run via plain Python."""
#     try:
#         import streamlit.runtime as _rt  # type: ignore
#         return bool(getattr(_rt, "exists", lambda: False)())
#     except Exception:
#         # Fallback env flags used by various Streamlit versions
#         return bool(
#             os.environ.get("STREAMLIT_SERVER_RUNNING")
#             or os.environ.get("STREAMLIT_RUNTIME")
#         )
#
# if _in_streamlit_runtime():
#     # Streamlit is running this script; render the app.
#     main_app()
# elif __name__ == "__main__":
#     # Prevent infinite relaunch. Tell the user how to start the app.
#     print("Start the app with:  streamlit run rig_app.py")
#
# print("Final notes ready.")


# rig_app.py
# ---------------------------------------------------------------------------
# Section: imports
# ---------------------------------------------------------------------------
from __future__ import annotations

import os
import sys
import json
import time
import math
import hashlib
import logging
import shutil
import re
from dataclasses import dataclass, asdict
from typing import List, Dict, Optional, Tuple, Iterable

import numpy as np
import faiss
import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI

# Document parsers
import fitz  # PyMuPDF
from docx import Document as DocxDocument

try:
    import textract  # optional for .doc
    TEXTRACT_AVAILABLE = True
except Exception:
    TEXTRACT_AVAILABLE = False

print("imports: ok")

# ---------------------------------------------------------------------------
# Section: user config
# ---------------------------------------------------------------------------
def _default_project_root() -> str:
    """Return default project root for Windows path."""
    return r"C:\Users\Andrew\PycharmProjects\Data Enablement Chatbot"


PROJECT_ROOT: str = os.getenv("PROJECT_ROOT", _default_project_root())
ENV_PATH: str = os.path.join(PROJECT_ROOT, ".env")

# Load .env first so env vars are available
load_dotenv(ENV_PATH)
load_dotenv()  # also load from current working dir if present

# Models and behavior
EMBEDDING_MODEL: str = "text-embedding-3-large"
CHAT_MODEL: str = "gpt-4o-mini"
EMBED_DIM: int = 3072
EMBED_BATCH_SIZE: int = 64

# Chunking
CHUNK_SIZE: int = 1800
CHUNK_OVERLAP: int = 250

# Retrieval
DEFAULT_TOP_K: int = 6

# Generation
DEFAULT_TEMPERATURE: float = 0.1
DEFAULT_MAX_TOKENS: int = 700

# Storage file names (inside index dir)
INDEX_DIR_NAME: str = ".rig_index"
FAISS_FILE: str = "faiss.index"
DOCSTORE_FILE: str = "docstore.jsonl"
MANIFEST_FILE: str = "manifest.json"
IDS_FILE: str = "ids.json"
EMBEDDINGS_FILE: str = "embeddings.npy"

# Docs folder from env; create if missing
DOCS_FOLDER: str = os.getenv(
    "DOCS_FOLDER",
    os.path.join(PROJECT_ROOT, "LocalDocs"),
)

os.makedirs(DOCS_FOLDER, exist_ok=True)

# Index dir sits inside project root to keep clean
INDEX_DIR: str = os.path.join(PROJECT_ROOT, INDEX_DIR_NAME)
os.makedirs(INDEX_DIR, exist_ok=True)

print("user config: ok")

# ---------------------------------------------------------------------------
# Section: logging setup
# ---------------------------------------------------------------------------
LOG_LEVEL = os.getenv("RIG_LOG_LEVEL", "INFO").upper()
logging.basicConfig(
    level=getattr(logging, LOG_LEVEL, logging.INFO),
    format="%(asctime)s | %(levelname)s | %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)],
)
logger = logging.getLogger("rig_app")
print("logging setup: ok")

# ---------------------------------------------------------------------------
# Section: data models
# ---------------------------------------------------------------------------
@dataclass
class Chunk:
    """A single text chunk with metadata."""
    id: int
    file_path: str
    file_name: str
    file_type: str
    page: Optional[int]  # 1-based for PDFs, else None
    chunk_index: int
    char_start: int
    char_end: int
    sha256: str
    text: str


@dataclass
class ManifestEntry:
    """Per-file manifest entry for incremental indexing."""
    file_path: str
    mtime: float
    sha256: str
    chunk_count: int
    file_type: str


print("data models: ok")

# ---------------------------------------------------------------------------
# Section: utilities
# ---------------------------------------------------------------------------
CONTROL_CHAR_PATTERN = re.compile(r"[\u0000-\u0008\u000B-\u001F\u007F]")
TIMESTAMP_PATTERN = re.compile(r"\[?\b\d{1,2}:\d{2}:\d{2}\b\]?|\[?\b\d{1,2}:\d{2}\b\]?")
SPEAKER_ARROW_PATTERN = re.compile(r"(?m)^\s*>+\s*")
REPEAT_PUNCT_PATTERN = re.compile(r"([.!?])\1{2,}")
WHITESPACE_RUN_PATTERN = re.compile(r"[ \t]{2,}")

ALLOWED_EXTS = {".pdf", ".docx", ".txt", ".doc"}


def ensure_dirs() -> None:
    """Ensure core directories exist."""
    os.makedirs(DOCS_FOLDER, exist_ok=True)
    os.makedirs(INDEX_DIR, exist_ok=True)


def sha256_file(path: str) -> str:
    """Compute SHA-256 of a file."""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def clean_text(text: str) -> str:
    """Clean text for robust transcripts."""
    if not text:
        return ""
    text = CONTROL_CHAR_PATTERN.sub("", text)
    text = TIMESTAMP_PATTERN.sub("", text)
    text = SPEAKER_ARROW_PATTERN.sub("", text)
    text = REPEAT_PUNCT_PATTERN.sub(r"\1", text)
    # Keep paragraph breaks, collapse extra spaces
    text = WHITESPACE_RUN_PATTERN.sub(" ", text)
    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse 3+ blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def paragraph_aware_chunks(text: str, chunk_size: int, overlap: int) -> List[Tuple[int, int, str]]:
    """Split into chunks ~chunk_size with overlap, respecting paragraphs."""
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    pos = 0
    chunks: List[Tuple[int, int, str]] = []
    buf = []
    buf_len = 0
    start_pos = 0

    for p in paragraphs:
        p_len = len(p) + 2  # account for paragraph break that was removed
        if buf_len + p_len <= chunk_size or not buf:
            if not buf:
                start_pos = pos
            buf.append(p)
            buf_len += p_len
            pos += p_len
        else:
            chunk_text = "\n\n".join(buf)
            end_pos = start_pos + len(chunk_text)
            chunks.append((start_pos, end_pos, chunk_text))

            # overlap by characters
            overlap_chars = min(overlap, len(chunk_text))
            window = chunk_text[max(0, len(chunk_text) - overlap_chars):]

            buf = [window, p] if window else [p]
            buf_len = len("\n\n".join(buf))
            start_pos = end_pos - overlap_chars
            pos += p_len

    if buf:
        chunk_text = "\n\n".join(buf)
        end_pos = start_pos + len(chunk_text)
        chunks.append((start_pos, end_pos, chunk_text))

    return chunks


def save_json(path: str, data: Dict) -> None:
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    os.replace(tmp, path)


def load_json(path: str) -> Dict:
    if not os.path.exists(path):
        return {}
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def save_jsonl(path: str, rows: Iterable[Dict]) -> None:
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        for row in rows:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")
    os.replace(tmp, path)


def load_jsonl(path: str) -> List[Dict]:
    if not os.path.exists(path):
        return []
    rows: List[Dict] = []
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            if line.strip():
                rows.append(json.loads(line))
    return rows


def citation_label(chunk: Chunk) -> str:
    """Return standardized citation label."""
    if chunk.file_type.lower() == ".pdf" and chunk.page:
        return f"[{chunk.file_name} p. {chunk.page}]"
    return f"[{chunk.file_name} chunk {chunk.chunk_index}]"


print("utilities: ok")

# ---------------------------------------------------------------------------
# Section: indexing
# ---------------------------------------------------------------------------
def scan_files(folder: str) -> List[str]:
    """Scan folder for allowed files."""
    paths: List[str] = []
    for root, _dirs, files in os.walk(folder):
        for name in files:
            ext = os.path.splitext(name)[1].lower()
            if ext in ALLOWED_EXTS:
                paths.append(os.path.join(root, name))
    return sorted(paths)


def parse_pdf(path: str) -> List[Tuple[str, Optional[int]]]:
    """Extract page-wise text from PDF."""
    out: List[Tuple[str, Optional[int]]] = []
    doc = fitz.open(path)
    try:
        for i, page in enumerate(doc, start=1):
            txt = page.get_text("text")
            txt = clean_text(txt)
            if txt:
                out.append((txt, i))
    finally:
        doc.close()
    return out


def parse_docx(path: str) -> str:
    """Extract text from DOCX, including tables."""
    doc = DocxDocument(path)
    parts: List[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return clean_text("\n\n".join(parts))


def parse_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return clean_text(f.read())


def parse_doc(path: str) -> str:
    if not TEXTRACT_AVAILABLE:
        logger.warning(".doc parsing requires textract. Skipping %s", path)
        return ""
    try:
        data = textract.process(path)
        return clean_text(data.decode("utf-8", errors="ignore"))
    except Exception as e:
        logger.error("textract failed for %s: %s", path, e)
        return ""


def extract_file(path: str) -> List[Chunk]:
    """Extract and chunk a file into Chunk objects."""
    ext = os.path.splitext(path)[1].lower()
    file_name = os.path.basename(path)
    file_sha = sha256_file(path)

    raw_segments: List[Tuple[str, Optional[int]]] = []
    if ext == ".pdf":
        raw_segments = parse_pdf(path)
    elif ext == ".docx":
        text = parse_docx(path)
        raw_segments = [(text, None)]
    elif ext == ".txt":
        text = parse_txt(path)
        raw_segments = [(text, None)]
    elif ext == ".doc":
        text = parse_doc(path)
        raw_segments = [(text, None)] if text else []
    else:
        return []

    chunks: List[Chunk] = []
    global_id_counter = 0  # temp id, will be reassigned later by build_index
    for seg_text, page in raw_segments:
        if not seg_text:
            continue
        spans = paragraph_aware_chunks(seg_text, CHUNK_SIZE, CHUNK_OVERLAP)
        for idx, (c_start, c_end, c_text) in enumerate(spans):
            chunks.append(
                Chunk(
                    id=-1,  # filled later
                    file_path=path,
                    file_name=file_name,
                    file_type=ext,
                    page=page,
                    chunk_index=len(chunks),  # sequential over file
                    char_start=c_start,
                    char_end=c_end,
                    sha256=file_sha,
                    text=c_text,
                )
            )
            global_id_counter += 1
    return chunks


def load_existing_index() -> Tuple[List[Chunk], Dict[str, ManifestEntry], np.ndarray]:
    """Load existing docstore, manifest, and embeddings if present."""
    manifest_raw = load_json(os.path.join(INDEX_DIR, MANIFEST_FILE))
    manifest: Dict[str, ManifestEntry] = {}
    for k, v in manifest_raw.items():
        manifest[k] = ManifestEntry(**v)

    doc_rows = load_jsonl(os.path.join(INDEX_DIR, DOCSTORE_FILE))
    chunks: List[Chunk] = [Chunk(**row) for row in doc_rows] if doc_rows else []

    emb_path = os.path.join(INDEX_DIR, EMBEDDINGS_FILE)
    embeddings = np.load(emb_path) if os.path.exists(emb_path) else np.empty((0, EMBED_DIM), dtype="float32")
    return chunks, manifest, embeddings


def embed_texts(client: OpenAI, texts: List[str]) -> np.ndarray:
    """Embed a list of texts with batching."""
    all_vecs: List[List[float]] = []
    for i in range(0, len(texts), EMBED_BATCH_SIZE):
        batch = texts[i : i + EMBED_BATCH_SIZE]
        resp = client.embeddings.create(model=EMBEDDING_MODEL, input=batch)
        vecs = [d.embedding for d in resp.data]
        all_vecs.extend(vecs)
    arr = np.array(all_vecs, dtype="float32")
    # L2 normalize for inner-product
    faiss.normalize_L2(arr)
    return arr


def build_faiss(embeddings: np.ndarray) -> faiss.IndexFlatIP:
    """Build FAISS inner-product index on normalized vectors."""
    index = faiss.IndexFlatIP(EMBED_DIM)
    if embeddings.size:
        index.add(embeddings)
    return index


def save_index(
    chunks: List[Chunk],
    manifest: Dict[str, ManifestEntry],
    embeddings: np.ndarray,
) -> None:
    """Persist index artifacts."""
    # Reassign contiguous ids
    for i, ch in enumerate(chunks):
        ch.id = i

    # Save docstore (without embeddings)
    save_jsonl(os.path.join(INDEX_DIR, DOCSTORE_FILE), [asdict(c) for c in chunks])

    # Save manifest
    save_json(
        os.path.join(INDEX_DIR, MANIFEST_FILE),
        {k: asdict(v) for k, v in manifest.items()},
    )

    # Save ids
    save_json(
        os.path.join(INDEX_DIR, IDS_FILE),
        {"ids": list(range(len(chunks)))},
    )

    # Save embeddings
    tmp = os.path.join(INDEX_DIR, EMBEDDINGS_FILE + ".tmp")
    np.save(tmp, embeddings)
    os.replace(tmp, os.path.join(INDEX_DIR, EMBEDDINGS_FILE))

    # Save FAISS
    faiss_path = os.path.join(INDEX_DIR, FAISS_FILE)
    faiss.write_index(build_faiss(embeddings), faiss_path)


def incremental_reindex(client: OpenAI, docs_folder: str) -> Tuple[int, int, float]:
    """Reindex incrementally. Returns (new_chunks, total_chunks, seconds)."""
    t0 = time.time()
    ensure_dirs()

    existing_chunks, existing_manifest, existing_embeddings = load_existing_index()
    path_to_chunks: Dict[str, List[Chunk]] = {}
    for ch in existing_chunks:
        path_to_chunks.setdefault(ch.file_path, []).append(ch)

    # Detect file states
    current_paths = scan_files(docs_folder)
    changed_paths: List[str] = []
    unchanged_paths: List[str] = []
    removed_paths: List[str] = []

    existing_paths_set = set(existing_manifest.keys())
    current_paths_set = set(current_paths)

    for path in current_paths:
        mtime = os.path.getmtime(path)
        sha = sha256_file(path)
        if path in existing_manifest:
            entry = existing_manifest[path]
            if not math.isclose(entry.mtime, mtime) or entry.sha256 != sha:
                changed_paths.append(path)
            else:
                unchanged_paths.append(path)
        else:
            changed_paths.append(path)

    removed_paths = list(existing_paths_set - current_paths_set)

    logger.info("Unchanged: %d | Changed/New: %d | Removed: %d", len(unchanged_paths), len(changed_paths), len(removed_paths))

    # Reuse unchanged chunks
    new_chunks: List[Chunk] = []
    for path in unchanged_paths:
        new_chunks.extend(path_to_chunks.get(path, []))

    # Extract changed files
    new_or_changed_chunks: List[Chunk] = []
    for path in changed_paths:
        cks = extract_file(path)
        new_or_changed_chunks.extend(cks)

    # Build new manifest
    manifest: Dict[str, ManifestEntry] = {}
    for path in unchanged_paths:
        entry = existing_manifest[path]
        manifest[path] = ManifestEntry(
            file_path=path,
            mtime=entry.mtime,
            sha256=entry.sha256,
            chunk_count=len(path_to_chunks.get(path, [])),
            file_type=os.path.splitext(path)[1].lower(),
        )
    for path in changed_paths:
        mtime = os.path.getmtime(path)
        sha = sha256_file(path)
        manifest[path] = ManifestEntry(
            file_path=path,
            mtime=mtime,
            sha256=sha,
            chunk_count=sum(1 for ch in new_or_changed_chunks if ch.file_path == path),
            file_type=os.path.splitext(path)[1].lower(),
        )

    # Merge chunks
    all_chunks: List[Chunk] = []
    all_chunks.extend(new_chunks)
    all_chunks.extend(new_or_changed_chunks)

    # Build embeddings with reuse for unchanged
    # Map file sha to embeddings for reuse
    reuse_embeds: Dict[int, np.ndarray] = {}
    if len(existing_chunks) == len(existing_embeddings):
        # safe mapping by id
        for ch in existing_chunks:
            if ch.file_path in unchanged_paths:
                reuse_embeds[ch.id] = existing_embeddings[ch.id : ch.id + 1]

    # Compose embeddings in new order
    embed_texts_list: List[str] = []
    embed_positions: List[int] = []
    collected_embeds: List[np.ndarray] = []

    for i, ch in enumerate(all_chunks):
        # cannot reuse by old id since ids will be reassigned
        # reuse based on identical file and matching text boundaries and sha
        reused = False
        if ch.file_path in unchanged_paths:
            # find a matching old chunk with same file, sha, char range
            for old in path_to_chunks.get(ch.file_path, []):
                if (old.sha256 == ch.sha256) and (old.char_start == ch.char_start) and (old.char_end == ch.char_end):
                    if old.id in reuse_embeds:
                        collected_embeds.append(reuse_embeds[old.id])
                        reused = True
                        break
        if not reused:
            embed_texts_list.append(ch.text)
            embed_positions.append(i)
            collected_embeds.append(None)  # placeholder

    # Compute new embeddings for changed chunks
    if embed_texts_list:
        new_embeds = embed_texts(client, embed_texts_list)
        # fill collected list in order
        n_idx = 0
        for j, arr in enumerate(collected_embeds):
            if arr is None:
                collected_embeds[j] = new_embeds[n_idx : n_idx + 1]
                n_idx += 1

    # Stack to final matrix
    embeddings = np.vstack(collected_embeds) if collected_embeds else np.empty((0, EMBED_DIM), dtype="float32")

    # Save everything
    save_index(all_chunks, manifest, embeddings)

    dt = time.time() - t0
    added = sum(1 for p in changed_paths)
    return len(new_or_changed_chunks), len(all_chunks), dt


def clear_index() -> None:
    """Delete index directory."""
    if os.path.exists(INDEX_DIR):
        shutil.rmtree(INDEX_DIR)
    os.makedirs(INDEX_DIR, exist_ok=True)
    logger.info("Index cleared.")


print("indexing: ok")

# ---------------------------------------------------------------------------
# Section: extraction
# ---------------------------------------------------------------------------
# (Parsing functions already included above in indexing)
print("extraction: ok")

# ---------------------------------------------------------------------------
# Section: OpenAI helpers
# ---------------------------------------------------------------------------
def get_openai_client() -> OpenAI:
    """Create OpenAI client from env var."""
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY not set. Add it to your .env file.")
    return OpenAI(api_key=api_key)


def chat_answer(
    client: OpenAI,
    question: str,
    retrieved: List[Chunk],
    temperature: float,
    max_tokens: int,
) -> str:
    """Ask the chat model constrained to snippets."""
    context_lines: List[str] = []
    for ch in retrieved:
        label = citation_label(ch)
        snippet = ch.text.strip().replace("\n", " ").strip()
        snippet = re.sub(r"\s{2,}", " ", snippet)
        context_lines.append(f"{label} :: {snippet}")

    system_msg = (
        "You are a strict RAG assistant. Use only the provided snippets. "
        "If the snippets are insufficient, say: \"I couldn’t find that in the provided documents.\" "
        "Always include inline citations for any facts you state in the form [filename.ext p. N] for PDFs "
        "or [filename.ext chunk K] for others. Keep answers clear and concise."
    )

    user_msg = (
        f"Question:\n{question}\n\n"
        "Snippets:\n" + "\n".join(context_lines)
    )

    resp = client.chat.completions.create(
        model=CHAT_MODEL,
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        temperature=float(temperature),
        max_tokens=int(max_tokens),
    )
    return resp.choices[0].message.content.strip()


print("OpenAI helpers: ok")

# ---------------------------------------------------------------------------
# Section: retrieval+answer
# ---------------------------------------------------------------------------
def load_runtime_index() -> Tuple[faiss.IndexFlatIP, List[Chunk], np.ndarray]:
    """Load FAISS, docstore, and embeddings for retrieval."""
    faiss_path = os.path.join(INDEX_DIR, FAISS_FILE)
    doc_rows = load_jsonl(os.path.join(INDEX_DIR, DOCSTORE_FILE))
    if not doc_rows or not os.path.exists(faiss_path):
        return faiss.IndexFlatIP(EMBED_DIM), [], np.empty((0, EMBED_DIM), dtype="float32")
    chunks = [Chunk(**row) for row in doc_rows]
    index = faiss.read_index(faiss_path)
    embeddings = np.load(os.path.join(INDEX_DIR, EMBEDDINGS_FILE))
    return index, chunks, embeddings


def retrieve(
    client: OpenAI,
    query: str,
    top_k: int = DEFAULT_TOP_K,
) -> List[Chunk]:
    """Retrieve top_k chunks."""
    index, chunks, _emb = load_runtime_index()
    if not chunks or index.ntotal == 0:
        return []

    q_vec = embed_texts(client, [query])
    D, I = index.search(q_vec, top_k)
    ids = I[0].tolist()
    out: List[Chunk] = []
    for idx in ids:
        if idx < 0 or idx >= len(chunks):
            continue
        out.append(chunks[idx])
    return out


def ensure_citations(answer: str, retrieved: List[Chunk]) -> str:
    """If no bracketed citations present, append Sources footer."""
    if re.search(r"\[[^\]]+\]", answer):
        return answer
    unique_labels = []
    for ch in retrieved:
        label = citation_label(ch)
        if label not in unique_labels:
            unique_labels.append(label)
    if not unique_labels:
        return 'I couldn’t find that in the provided documents.'
    footer = "\n\nSources: " + "; ".join(unique_labels)
    return answer + footer


def format_top_sources(chunks: List[Chunk], limit: int = 6) -> List[str]:
    """Return preview lines for UI."""
    lines: List[str] = []
    for ch in chunks[:limit]:
        preview = ch.text.strip().split("\n", 1)[0]
        preview = re.sub(r"\s{2,}", " ", preview)
        lines.append(f"{citation_label(ch)} — {preview[:200]}...")
    return lines


print("retrieval+answer: ok")

# ---------------------------------------------------------------------------
# Section: Streamlit UI
# ---------------------------------------------------------------------------
def write_requirements_file(project_root: str) -> str:
    """Write requirements_rig.txt and build_win.bat to project root."""
    req_path = os.path.join(project_root, "requirements_rig.txt")
    bat_path = os.path.join(project_root, "build_win.bat")

    requirements_text = (
        "faiss-cpu==1.8.0.post1\n"
        "numpy==1.26.4\n"
        "PyMuPDF==1.24.9\n"
        "streamlit==1.37.1\n"
        "openai==1.42.0\n"
        "python-docx==1.1.2\n"
        "tiktoken==0.7.0\n"
        "textract==1.6.5\n"
        "python-dotenv==1.0.1\n"
    )
    with open(req_path, "w", encoding="utf-8") as f:
        f.write(requirements_text)

    bat_text = (
        "@echo off\n"
        "setlocal enableextensions\n"
        "cd /d %~dp0\n"
        "IF NOT EXIST .venv (\n"
        "  py -m venv .venv\n"
        ")\n"
        "call .venv\\Scripts\\activate\n"
        "pip install --upgrade pip\n"
        "pip install -r requirements_rig.txt\n"
        "pip install pyinstaller\n"
        "pyinstaller --onefile --windowed --name rig_app rig_app.py\n"
        "echo Built EXE in .\\dist\\\n"
        "echo Note: The EXE reads OPENAI_API_KEY and DOCS_FOLDER from a .env file in the project root.\n"
        "endlocal\n"
    )
    with open(bat_path, "w", encoding="utf-8") as f:
        f.write(bat_text)

    return f"Wrote {req_path} and {bat_path}"


def sidebar_controls() -> Dict:
    """Render sidebar and return config dict."""
    st.sidebar.header("RIG Controls")
    st.sidebar.caption("Local, docs-only RAG")

    # Docs folder from env, allow override for testing
    current_docs = st.sidebar.text_input(
        "Documents folder",
        value=DOCS_FOLDER,
        help="Path is also read from DOCS_FOLDER in .env",
    )
    if current_docs and current_docs != DOCS_FOLDER:
        # update globals for this session
        st.session_state["docs_folder"] = current_docs
        os.makedirs(current_docs, exist_ok=True)
    docs_folder = st.session_state.get("docs_folder", DOCS_FOLDER)

    # File uploader to add to local KB
    uploaded = st.sidebar.file_uploader(
        "Add documents",
        type=["pdf", "docx", "txt", "doc"],
        accept_multiple_files=True,
        help="Dropped files are saved into the Documents folder",
    )
    if uploaded:
        for up in uploaded:
            ext = os.path.splitext(up.name)[1].lower()
            if ext not in ALLOWED_EXTS:
                st.sidebar.warning(f"Skipped {up.name} (unsupported type)")
                continue
            save_path = os.path.join(docs_folder, up.name)
            try:
                with open(save_path, "wb") as f:
                    f.write(up.getbuffer())
                st.sidebar.success(f"Saved {up.name}")
            except Exception as e:
                st.sidebar.error(f"Failed to save {up.name}: {e}")

    # Testing options
    st.sidebar.subheader("Testing options")
    top_k = st.sidebar.slider("Top-k", 1, 12, DEFAULT_TOP_K)
    temperature = st.sidebar.slider("Temperature", 0.0, 1.0, DEFAULT_TEMPERATURE, 0.05)
    max_tokens = st.sidebar.slider("Max tokens", 200, 2000, DEFAULT_MAX_TOKENS, 50)
    show_debug = st.sidebar.checkbox("Show retrieval previews", value=True)
    force_reindex = st.sidebar.button("Reindex")
    if st.sidebar.button("Clear index"):
        clear_index()
        st.sidebar.success("Index cleared.")

    if st.sidebar.button("Write requirements"):
        msg = write_requirements_file(PROJECT_ROOT)
        st.sidebar.success(msg)

    if force_reindex:
        try:
            client = get_openai_client()
            added, total, secs = incremental_reindex(client, docs_folder)
            st.sidebar.success(f"Reindexed. New/changed chunks: {added}. Total chunks: {total}. Time: {secs:.1f}s")
        except Exception as e:
            st.sidebar.error(f"Reindex failed: {e}")

    return {
        "docs_folder": docs_folder,
        "top_k": top_k,
        "temperature": temperature,
        "max_tokens": max_tokens,
        "show_debug": show_debug,
    }


def run_streamlit_app() -> None:
    """Main Streamlit app."""
    st.set_page_config(page_title="Docs-only RIG", layout="wide")
    st.title("Docs-only RAG Chatbot")
    st.caption("Uses only your local folder as the source of truth.")

    cfg = sidebar_controls()

    # Check API key
    api_present = bool(os.getenv("OPENAI_API_KEY", "").strip())
    if not api_present:
        st.error("OPENAI_API_KEY is not set. Add it to your .env file in the project root and reload.")
        st.stop()

    # Ensure index exists
    if not os.path.exists(os.path.join(INDEX_DIR, FAISS_FILE)):
        with st.spinner("Building index for the first time..."):
            try:
                client = get_openai_client()
                added, total, secs = incremental_reindex(client, cfg["docs_folder"])
                st.success(f"Indexed. New chunks: {added}. Total: {total}. Time: {secs:.1f}s")
            except Exception as e:
                st.error(f"Initial index failed: {e}")
                st.stop()

    # Chat history
    if "history" not in st.session_state:
        st.session_state["history"] = []

    with st.form("qa_form", clear_on_submit=False):
        question = st.text_area("Ask a question about your documents", height=100)
        submitted = st.form_submit_button("Ask")

    if st.button("Clear chat"):
        st.session_state["history"] = []
        st.rerun()

    if submitted and question.strip():
        try:
            client = get_openai_client()
            retrieved = retrieve(client, question, top_k=cfg["top_k"])
            if not retrieved:
                answer = 'I couldn’t find that in the provided documents.'
            else:
                raw_answer = chat_answer(
                    client=client,
                    question=question,
                    retrieved=retrieved,
                    temperature=cfg["temperature"],
                    max_tokens=cfg["max_tokens"],
                )
                answer = ensure_citations(raw_answer, retrieved)

            st.session_state["history"].append(
                {"question": question.strip(), "answer": answer, "retrieved": retrieved}
            )
        except Exception as e:
            st.error(f"Query failed: {e}")

    # Render history
    for turn in st.session_state["history"]:
        st.markdown(f"**You:** {turn['question']}")
        st.markdown(f"**Answer:** {turn['answer']}")
        if cfg["show_debug"]:
            with st.expander("Top sources"):
                previews = format_top_sources(turn["retrieved"])
                for line in previews:
                    st.write(line)
        st.markdown("---")


print("Streamlit UI: ok")

# ---------------------------------------------------------------------------
# Section: requirements snippet
# ---------------------------------------------------------------------------
# Written by the UI button into project root. See write_requirements_file().
print("requirements snippet: ok")

# ---------------------------------------------------------------------------
# Section: final notes
# ---------------------------------------------------------------------------
def main() -> None:
    """Entrypoint. Run streamlit when executed as a script or packaged."""
    # When run directly, hand off to streamlit runner so it behaves the same.
    if os.environ.get("STREAMLIT_SERVER_ENABLED", ""):
        # Already under Streamlit
        run_streamlit_app()
        return
    # Launch Streamlit programmatically
    try:
        import streamlit.web.cli as stcli
        sys.argv = ["streamlit", "run", os.path.abspath(__file__), "--server.headless=true"]
        stcli.main()
    except SystemExit:
        # Streamlit may call sys.exit; ignore to allow PyInstaller to exit cleanly
        pass


if __name__ == "__main__":
    main()

print("final notes: ok")
